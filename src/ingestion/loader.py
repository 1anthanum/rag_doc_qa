"""
Document loader supporting multiple file formats.
Extracts raw text from PDF, TXT, Markdown, and DOCX files.
"""

import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import List

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a loaded document with metadata."""

    content: str
    source: str
    doc_type: str
    metadata: dict = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.content)

    @property
    def word_count(self) -> int:
        return len(self.content.split())


class DocumentLoader:
    """
    Unified document loader that dispatches to format-specific parsers.

    Supported formats:
        - .txt / .md  : Plain text reading
        - .pdf        : PyPDF2-based extraction
        - .docx       : python-docx extraction

    Usage:
        loader = DocumentLoader()
        docs = loader.load_directory("./data/documents/")
        doc = loader.load_file("paper.pdf")
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

    def load_file(self, file_path: str) -> Document:
        """Load a single file and return a Document object."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        logger.info(f"Loading {ext} file: {path.name}")

        if ext in (".txt", ".md"):
            content = self._load_text(path)
        elif ext == ".pdf":
            content = self._load_pdf(path)
        elif ext == ".docx":
            content = self._load_docx(path)
        else:
            raise ValueError(f"No loader for extension: {ext}")

        return Document(
            content=content,
            source=str(path.resolve()),
            doc_type=ext.lstrip("."),
            metadata={
                "filename": path.name,
                "size_bytes": path.stat().st_size,
            },
        )

    def load_directory(self, dir_path: str, recursive: bool = True) -> List[Document]:
        """Load all supported files from a directory."""
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in sorted(dir_path.glob(pattern)):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                    logger.info(
                        f"  Loaded: {file_path.name} " f"({doc.word_count} words)"
                    )
                except Exception as e:
                    logger.warning(f"  Failed to load {file_path.name}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {dir_path}")
        return documents

    # ── Format-specific parsers ──────────────────────────────────

    @staticmethod
    def _load_text(path: Path) -> str:
        """Load plain text or markdown file."""
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _load_pdf(path: Path) -> str:
        """Extract text from PDF using PyPDF2."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF loading. "
                "Install with: pip install PyPDF2"
            )

        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text)

        return "\n\n".join(pages)

    @staticmethod
    def _load_docx(path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install with: pip install python-docx"
            )

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
